"""
Generate Camilo Vélez's CV tailored for a Data Engineer role.
Reuses formatting helpers from the parent generate_cv module.
"""

import sys
import os

# Add parent to path for imports
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from generate_cv import (
    add_bottom_border,
    add_hyperlink,
    add_section_heading as _add_section_heading,
    add_job as _add_job,
    set_cell_border,
)


def add_section_heading(doc, title):
    """Tighter section heading for 1-page fit."""
    p = doc.add_paragraph()
    p.space_before = Pt(5)
    p.space_after = Pt(2)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    add_bottom_border(p)


def add_job(doc, title, dates, subtitle=None, bullets=None):
    """Tighter job entry for 1-page fit."""
    p = doc.add_paragraph()
    p.space_after = Pt(0)
    p.space_before = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(9.5)
    run_date = p.add_run(f"\t{dates}")
    run_date.italic = True
    run_date.font.size = Pt(9.5)

    if subtitle:
        p = doc.add_paragraph()
        p.space_after = Pt(0)
        run = p.add_run(subtitle)
        run.italic = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if bullets:
        for bullet in bullets:
            p = doc.add_paragraph(style="List Bullet")
            p.space_after = Pt(0)
            p.clear()
            run = p.add_run(bullet)
            run.font.size = Pt(9.5)


def create_cv():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(9.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # ── HEADER ──
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.space_after = Pt(0)
    run = name.add_run("Camilo Alejandro Vélez Medina")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(2)

    contact_items = [
        ("Brisbane, QLD", None),
        ("camilo95022@gmail.com", "mailto:camilo95022@gmail.com"),
        ("+61 450 438 679", None),
        ("LinkedIn", "https://www.linkedin.com/in/camilo-alejandro-v%C3%A9lez-medina-1b447a144"),
        ("GitHub", "https://github.com/Camilovelez1"),
    ]
    for i, (text, url) in enumerate(contact_items):
        if i > 0:
            sep = contact.add_run("  |  ")
            sep.font.size = Pt(9)
            sep.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        if url:
            add_hyperlink(contact, text, url)
        else:
            run = contact.add_run(text)
            run.font.size = Pt(9)

    # ── PROFILE ──
    add_section_heading(doc, "Profile")
    p = doc.add_paragraph()
    p.space_after = Pt(2)
    run = p.add_run(
        "Data Engineer and Tech Lead with a Master's in Data Science and 6+ years delivering "
        "production-grade data pipelines and analytics-ready data layers across government, healthcare, "
        "finance, and logistics sectors. Deep hands-on expertise in Databricks (PySpark & Spark SQL), "
        "dbt, Delta Lake medallion architecture, and Power BI semantic models. "
        "Proven track record designing optimised star schemas, implementing CI/CD for data pipelines, "
        "and enabling cross-functional teams through modern lakehouse environments on Azure."
    )
    run.font.size = Pt(9.5)

    # ── PROFESSIONAL EXPERIENCE ──
    add_section_heading(doc, "Professional Experience")

    # BPT
    add_job(
        doc,
        title="Tech Lead / Data Engineering – BPT",
        dates="2024 – present",
        subtitle="Government Tax Authority (SUNAT Peru), Pension Funds, Healthcare, Oil & Gas, Banking",
        bullets=[
            "Designed and maintained production-grade data pipelines using Databricks (PySpark & Spark SQL), ingesting data from multiple source systems with Delta Lake medallion architecture (bronze/silver/gold layers).",
            "Developed dbt models across staging, intermediate, and presentation layers; implemented dbt tests and data quality monitoring to ensure reliability and consistency.",
            "Built data models optimised for Power BI semantic models with star schemas; enabled Power Apps integration with Power BI datasets for operational applications.",
            "Conducted data gap analysis identifying missing tables, attributes, and transformations; produced technical documentation covering data sources and business logic.",
            "Managed Git-based version control and CI/CD for data pipelines; optimised Databricks workloads achieving 40% improvement in ETL processing times via Delta Live Tables.",
        ],
    )

    # DeltaWits
    add_job(
        doc,
        title="Data Engineer & Consultant – DeltaWits",
        dates="2023 – present",
        subtitle="Healthcare Technology (Remote)",
        bullets=[
            "Designed and maintained data pipelines for healthcare data processing on AWS, ingesting and transforming structured and unstructured data from multiple source systems into analytics-ready layers.",
            "Built scalable, compliant cloud infrastructure for high-volume data processing, implementing data quality checks and monitoring mechanisms across the pipeline.",
        ],
    )

    # Botero Soto - Data Scientist
    add_job(
        doc,
        title="Data Engineer / Data Scientist – Botero Soto",
        dates="2023 – 2024",
        bullets=[
            "Built and optimised data pipelines using Azure Databricks, designing analytics-ready data layers and Power BI dashboards with star schema models that achieved 30% response time improvements.",
        ],
    )

    # Botero Soto - Project Leader
    add_job(
        doc,
        title="Project Leader – Botero Soto",
        dates="2018 – 2023",
        bullets=[
            "Managed data-driven optimisation projects using agile methodologies; built Power BI dashboards and traceability systems. Supervised teams of up to 45 people; recognised for best operations performance (2020).",
        ],
    )

    # ── EDUCATION ──
    add_section_heading(doc, "Education")
    edu_items = [
        ("Master's in Data Science and Analytics", "EAFIT University (Top Colombian University)", "2023–2025"),
        ("Bachelor's in Business Administration", "UNIMINUTO University", "2013–2018"),
    ]
    for degree, school, years in edu_items:
        p = doc.add_paragraph()
        p.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        run = p.add_run(f"{degree}")
        run.bold = True
        run.font.size = Pt(9.5)
        p.add_run(f" – {school}").font.size = Pt(9.5)
        run_date = p.add_run(f"\t{years}")
        run_date.font.size = Pt(9.5)
        run_date.italic = True

    # ── TECHNICAL SKILLS ──
    add_section_heading(doc, "Technical Skills")
    skills = [
        ("Data Engineering", "Databricks, Spark SQL, PySpark, dbt (Data Build Tools), Delta Lake, Medallion Architecture, Delta Live Tables, Unity Catalog"),
        ("BI & Apps", "Power BI (Semantic Models, Star Schemas, DAX), Power Apps, Power Automate"),
        ("Cloud & CI/CD", "Azure (Data Factory, Synapse), AWS, CI/CD Pipelines, Git, GitHub"),
        ("Programming", "Python, SQL (Advanced), Pandas, NumPy, MLflow, Data Governance"),
        ("Leadership", "Team Management (45+ people), Agile Methodologies, Cross-functional Collaboration, Technical Documentation"),
    ]
    for category, items in skills:
        p = doc.add_paragraph()
        p.space_after = Pt(0)
        run = p.add_run(f"{category}: ")
        run.bold = True
        run.font.size = Pt(9.5)
        run = p.add_run(items)
        run.font.size = Pt(9.5)

    # ── CERTIFICATIONS ──
    add_section_heading(doc, "Certifications")

    # Create a table with 1 row, 2 columns for badge + text layout
    cert_table = doc.add_table(rows=1, cols=2)
    cert_table.autofit = False

    # Force grid column widths: narrow badge + wide text
    tbl = cert_table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is not None:
        tblW.set(qn("w:type"), "dxa")
        tblW.set(qn("w:w"), "9638")
    tblGrid = tbl.find(qn("w:tblGrid"))
    gridCols = tblGrid.findall(qn("w:gridCol"))
    gridCols[0].set(qn("w:w"), "1134")  # ~2cm for badge
    gridCols[1].set(qn("w:w"), "8504")  # rest for text

    # Badge cell (left, narrow)
    badge_cell = cert_table.cell(0, 0)
    badge_cell.width = Cm(2)
    badge_p = badge_cell.paragraphs[0]
    badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge_img_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", "..", "assets", "databricks_badge.png")
    if os.path.exists(badge_img_path):
        badge_run = badge_p.add_run()
        badge_run.add_picture(badge_img_path, width=Cm(1.3))

    # Certs text cell (right, wide)
    text_cell = cert_table.cell(0, 1)
    text_cell.width = Cm(15)
    # Remove default empty paragraph
    text_cell.paragraphs[0].clear()
    certs = [
        "Data Engineer Associate – Databricks (Certified) | Data Engineer Professional – Databricks (In Progress)",
        "Data Science in Python / SQL for Data Science – Coursera | Foundations of Data Science – EAFIT",
    ]
    for i, cert in enumerate(certs):
        if i == 0:
            p = text_cell.paragraphs[0]
        else:
            p = text_cell.add_paragraph()
        p.style = doc.styles["List Bullet"]
        p.space_after = Pt(1)
        p.clear()
        run = p.add_run(cert)
        run.font.size = Pt(9.5)

    # Remove table borders
    for row in cert_table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "none", "sz": "0", "color": "FFFFFF"},
                bottom={"val": "none", "sz": "0", "color": "FFFFFF"},
                start={"val": "none", "sz": "0", "color": "FFFFFF"},
                end={"val": "none", "sz": "0", "color": "FFFFFF"},
            )

    # ── FOOTER ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(4)
    run = p.add_run("References available upon request | Updated February 2026")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    # Save
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, "CV_Camilo_Velez_DataEngineer.docx")
    doc.save(output_path)
    print(f"CV saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_cv()
