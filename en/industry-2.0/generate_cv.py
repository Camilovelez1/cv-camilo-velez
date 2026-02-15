"""
Generate Camilo Vélez's CV as a professional Word document.
Content based on CV_Camilo_Velez_TMR_Optimized.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        if edge in kwargs:
            element = OxmlElement(f"w:{edge}")
            for attr, val in kwargs[edge].items():
                element.set(qn(f"w:{attr}"), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def add_bottom_border(paragraph):
    """Add a bottom border line under a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2B579A")
    pBdr.append(bottom)
    pPr.append(pBdr)


def create_cv():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.3)
        section.bottom_margin = Cm(1.3)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    # ── HEADER ──
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.space_after = Pt(2)
    run = name.add_run("Camilo Alejandro Vélez Medina")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(6)

    contact_items = [
        ("Brisbane, QLD", None),
        ("camilo95022@gmail.com", "mailto:camilo95022@gmail.com"),
        ("+61 450 438 679", None),
        ("LinkedIn", "https://www.linkedin.com/in/camilo-alejandro-v%C3%A9lez-medina-1b447a144"),
        ("GitHub", "https://github.com/Camilovelez1"),
    ]
    for i, (text, url) in enumerate(contact_items):
        if i > 0:
            sep = contact.add_run("  |  ")
            sep.font.size = Pt(9)
            sep.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        if url:
            add_hyperlink(contact, text, url)
        else:
            run = contact.add_run(text)
            run.font.size = Pt(9)

    # ── PROFILE ──
    add_section_heading(doc, "Profile")
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    run = p.add_run(
        "AI Data Engineer and Tech Lead with Master's in Data Science and 6+ years delivering "
        "scalable data solutions across government, healthcare, finance, and logistics sectors. "
        "Deep expertise in Azure Data Factory, Databricks, Snowflake, and AWS platforms, "
        "specializing in Python, PySpark, SQL, Machine Learning, and MLOps. "
        "Proven leader of cross-functional teams, consistently achieving 40%+ improvements "
        "in data processing performance while driving technical innovation in fast-paced environments."
    )
    run.font.size = Pt(10)

    # ── PROFESSIONAL EXPERIENCE ──
    add_section_heading(doc, "Professional Experience")

    # BPT
    add_job(
        doc,
        title="Tech Lead / Data Engineering – BPT",
        dates="2024 – present",
        subtitle="Government Tax Authority (SUNAT Peru), Pension Funds, Healthcare, Oil & Gas, Banking",
        bullets=[
            "Leading real-time data pipeline implementation for SUNAT (Peru's National Tax Administration), processing large-scale financial transaction data with strict government compliance requirements including CDC and streaming architectures.",
            "Architected scalable data solutions using Snowflake data warehousing and DBT (Data Build Tools), improving performance by 40% through optimized ETL processes and Delta Live Tables implementations.",
            "Developed knowledge graph implementations using Neo4j for enhanced data relationships and integrated LLM capabilities using OpenAI and Claude APIs for natural language query interfaces.",
            "Built interactive conversational interfaces leveraging Databricks Genie, democratizing data access for business users and enabling teams to query complex datasets using natural language.",
            "Managed international client relationships (95% satisfaction) while leading technical architecture decisions and maintaining hands-on involvement in code reviews.",
        ],
    )

    # DeltaWits
    add_job(
        doc,
        title="AI Engineer & Data Consultant – DeltaWits",
        dates="2023 – present",
        subtitle="Healthcare Technology & AI Solutions (Remote)",
        bullets=[
            "Architected AI-powered healthcare solutions using AWS Bedrock and LLMs, reducing physician preparation time by 40% through automated medical data processing and intelligent profile generation.",
            "Developed scalable, HIPAA-compliant cloud infrastructure on AWS for high-volume medical data processing, integrating structured and unstructured health data sources.",
        ],
    )

    # Botero Soto - Data Scientist
    add_job(
        doc,
        title="Data Scientist – Botero Soto",
        dates="2023 – 2024",
        bullets=[
            "Led predictive modeling using Azure Databricks achieving 30% response time improvements; designed scalable analytics solutions and automated processes with Power BI dashboards.",
        ],
    )

    # Botero Soto - Project Leader
    add_job(
        doc,
        title="Project Leader – Botero Soto",
        dates="2018 – 2023",
        bullets=[
            "Managed optimization projects using agile methodologies; supervised teams of up to 45 people and implemented high-efficiency programs recognized for best operations performance (2020).",
        ],
    )

    # ── EDUCATION ──
    add_section_heading(doc, "Education")
    edu_items = [
        ("Master's in Data Science and Analytics", "EAFIT University (Top Colombian University)", "2023–2025"),
        ("Bachelor's in Business Administration", "UNIMINUTO University", "2013–2018"),
    ]
    for degree, school, years in edu_items:
        p = doc.add_paragraph()
        p.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        run = p.add_run(f"{degree}")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(f" – {school}").font.size = Pt(10)
        run_date = p.add_run(f"\t{years}")
        run_date.font.size = Pt(10)
        run_date.italic = True

    # ── TECHNICAL SKILLS ──
    add_section_heading(doc, "Technical Skills")
    skills = [
        ("Data Platforms & Tools", "Snowflake, Databricks, Azure Data Factory, DBT (Data Build Tools), Unity Catalog, Delta Live Tables"),
        ("Programming & Analytics", "Python, PySpark, SQL, Pandas, NumPy, Power BI, MLflow, GitHub"),
        ("AI & Machine Learning", "LLMs, OpenAI API, Claude API, AWS Bedrock, MLOps, Knowledge Graphs, Neo4j"),
        ("Cloud Platforms", "Azure, AWS, Real-time Streaming, CDC (Change Data Capture), Data Governance"),
        ("Leadership & Collaboration", "Team Management (45+ people), Agile Methodologies, Cross-functional Collaboration"),
    ]
    for category, items in skills:
        p = doc.add_paragraph()
        p.space_after = Pt(1)
        run = p.add_run(f"{category}: ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(items)
        run.font.size = Pt(10)

    # ── CERTIFICATIONS ──
    add_section_heading(doc, "Certifications")

    # Create a table with 1 row, 2 columns for badge + text layout
    cert_table = doc.add_table(rows=1, cols=2)
    cert_table.autofit = True

    # Badge cell (left, narrow)
    badge_cell = cert_table.cell(0, 0)
    badge_cell.width = Cm(2)
    badge_p = badge_cell.paragraphs[0]
    badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge_img_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", "assets", "databricks_badge.png")
    if os.path.exists(badge_img_path):
        badge_run = badge_p.add_run()
        badge_run.add_picture(badge_img_path, width=Cm(1.8))

    # Certs text cell (right, wide)
    text_cell = cert_table.cell(0, 1)
    text_cell.width = Cm(15)
    # Remove default empty paragraph
    text_cell.paragraphs[0].clear()
    certs = [
        "Data Engineer Associate – Databricks (Certified) | Data Engineer Professional – Databricks (In Progress)",
        "Data Science in Python / SQL for Data Science – Coursera | Foundations of Data Science – EAFIT",
    ]
    for i, cert in enumerate(certs):
        if i == 0:
            p = text_cell.paragraphs[0]
        else:
            p = text_cell.add_paragraph()
        p.style = doc.styles["List Bullet"]
        p.space_after = Pt(1)
        p.clear()
        run = p.add_run(cert)
        run.font.size = Pt(10)

    # Remove table borders
    for row in cert_table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "none", "sz": "0", "color": "FFFFFF"},
                bottom={"val": "none", "sz": "0", "color": "FFFFFF"},
                start={"val": "none", "sz": "0", "color": "FFFFFF"},
                end={"val": "none", "sz": "0", "color": "FFFFFF"},
            )

    # ── FOOTER ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    run = p.add_run("References available upon request | Updated February 2026")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    # Save
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, "CV_Camilo_Velez.docx")
    doc.save(output_path)
    print(f"CV saved to: {output_path}")
    return output_path


def add_section_heading(doc, title):
    """Add a styled section heading with bottom border."""
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(4)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    add_bottom_border(p)


def add_job(doc, title, dates, subtitle=None, bullets=None):
    """Add a job entry."""
    # Title line with date
    p = doc.add_paragraph()
    p.space_after = Pt(1)
    p.space_before = Pt(4)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run_date = p.add_run(f"\t{dates}")
    run_date.italic = True
    run_date.font.size = Pt(10)

    if subtitle:
        p = doc.add_paragraph()
        p.space_after = Pt(1)
        run = p.add_run(subtitle)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if bullets:
        for bullet in bullets:
            p = doc.add_paragraph(style="List Bullet")
            p.space_after = Pt(1)
            p.clear()
            run = p.add_run(bullet)
            run.font.size = Pt(10)


def add_hyperlink(paragraph, text, url):
    """Add a real clickable hyperlink to a paragraph.
    Returns the hyperlink run so caller can style it.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Blue color
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2B579A")
    rPr.append(color)

    # Underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    # Font size
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")  # half-points, 18 = 9pt
    rPr.append(sz)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return new_run


if __name__ == "__main__":
    create_cv()
